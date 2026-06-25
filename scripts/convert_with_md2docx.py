#!/usr/bin/env python3
"""
Convert markdown to DOCX using md2docx-python library with custom header formatting.

This script provides an alternative conversion method that may handle tables better
than pandoc, while still applying our custom header and classification formatting.

Usage:
    python convert_with_md2docx.py <input.md> <output.docx> [title] [classification]

Dependencies:
    pip install md2docx-python python-docx
    
    Note: The md2docx-python package is imported as:
    from md2docx_python.src.md2docx_python import markdown_to_word
"""

import sys
import os

# Fix for GitHub Actions: add user site-packages to path
if sys.platform.startswith('linux'):
    # GitHub Actions uses ~/.local/lib/python3.X/site-packages
    user_site = os.path.expanduser('~/.local/lib/python3.12/site-packages')
    if os.path.exists(user_site) and user_site not in sys.path:
        sys.path.insert(0, user_site)
    
    # Also try common fallback paths
    for version in ['3.11', '3.10', '3.9']:
        fallback_path = os.path.expanduser(f'~/.local/lib/python{version}/site-packages')
        if os.path.exists(fallback_path) and fallback_path not in sys.path:
            sys.path.insert(0, fallback_path)
    
    # Try using site module as well
    try:
        import site
        user_site = site.getusersitepackages()
        if user_site and user_site not in sys.path:
            sys.path.insert(0, user_site)
    except:
        pass

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def update_docx_header(docx_path, title_text, classification=None, font_size=14, bold=True, align='center'):
    """
    Update header in DOCX file with title and classification.
    Classification appears on the right side of the header.
    """
    docx_path = os.path.abspath(docx_path)
    
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Error: Cannot find DOCX file at '{docx_path}'")
    
    doc = Document(docx_path)
    
    # Map alignment strings to constants
    alignment_map = {
        'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
        'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
        'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
    }
    alignment = alignment_map.get(align.lower(), WD_PARAGRAPH_ALIGNMENT.CENTER)
    
    # Update the headers in all sections
    for section in doc.sections:
        header = section.header
        
        # Get all paragraphs first to avoid iteration issues
        paragraphs = list(header.paragraphs)
        
        # Clear all existing paragraphs by setting their text to empty
        for paragraph in paragraphs:
            paragraph.text = ''
        
        # Use or create first paragraph for the title
        if paragraphs:
            header_para = paragraphs[0]
        else:
            header_para = header.add_paragraph()
        
        # Set title text (preserve original case)
        header_para.text = title_text
        
        # Apply styling to title
        for run in header_para.runs:
            run.font.size = Pt(font_size)
            run.font.bold = bold
        header_para.alignment = alignment
        
        # Add classification text on the RIGHT side of the same line if provided
        if classification:
            # Create a single line: Title                              UNCLASSIFIED
            # We'll use tabs to separate them for right alignment
            header_para.text = f"{title_text}\t\t{classification}"
            
            # Style the entire line
            for run in header_para.runs:
                run.font.size = Pt(font_size)
                run.font.bold = bold
    
    # Save the updated document
    doc.save(docx_path)
    print(f"✅ Updated header in '{docx_path}'")


def convert_with_md2docx(input_md, output_docx, title=None, classification=None):
    """
    Convert markdown to DOCX using md2docx-python library.
    
    Args:
        input_md: Path to input markdown file
        output_docx: Path to output DOCX file
        title: Document title (optional)
        classification: Classification text (optional)
    """
    try:
        # Correct import path for md2docx-python
        from md2docx_python.src.md2docx_python import markdown_to_word as convert
        
        print(f"📄 Converting '{input_md}' to '{output_docx}' using md2docx...")
        
        # Convert markdown to DOCX
        convert(input_md, output_docx)
        
        # Apply custom header if title is provided
        if title:
            update_docx_header(output_docx, title, classification)
        
        print(f"✅ Conversion complete!")
        return True
        
    except ImportError:
        print("❌ Error: md2docx-python is not installed.")
        print("   Install it with: pip install md2docx-python")
        return False
    except Exception as e:
        print(f"❌ Error during conversion: {e}")
        return False


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python convert_with_md2docx.py <input.md> <output.docx> [title] [classification]")
        print()
        print("Example: python convert_with_md2docx.py docs/ir-05.md output/ir-05.docx \\")
        print("             'Aurora IR-4 Container Breach' 'UNCLASSIFIED'")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    title = sys.argv[3] if len(sys.argv) > 3 else None
    classification = sys.argv[4] if len(sys.argv) > 4 else None
    
    success = convert_with_md2docx(input_file, output_file, title, classification)
    sys.exit(0 if success else 1)
