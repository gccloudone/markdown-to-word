import sys
from docx import Document

def update_header(docx_path, title_text):
    doc = Document(docx_path)
    # Loop through every section (Header, Footer, etc.)
    for section in doc.sections:
        header = section.header
        # Replace the placeholder text or set new text
        if not header.paragraphs:
            header.add_paragraph(title_text)
        else:
            header.paragraphs[0].text = title_text
    doc.save(docx_path)

if __name__ == "__main__":
    update_header(sys.argv[1], sys.argv[2])
