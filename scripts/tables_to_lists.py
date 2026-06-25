import sys
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def table_to_list_text(table, list_style='bullet', separator=': '):
    """
    Convert a table to a list of text strings.
    
    Args:
        table: The table object to convert
        list_style: 'bullet', 'number', or 'dash' (default: 'bullet')
        separator: String to separate column headers from values (default: ': ')
    
    Returns:
        List of text strings representing the table data as list items
    """
    text_items = []
    
    # Get table data
    table_data = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        table_data.append(row_data)
    
    if not table_data:
        return text_items
    
    # Determine if first row is headers
    is_header_row = False
    if len(table_data) > 1:
        # Check if first row cells are shorter (likely headers)
        avg_first_row = sum(len(cell) for cell in table_data[0]) / max(len(table_data[0]), 1)
        total_other_chars = sum(len(cell) for row in table_data[1:] for cell in row)
        total_other_cells = sum(len(row) for row in table_data[1:])
        avg_other_rows = total_other_chars / max(total_other_cells, 1)
        if avg_first_row < avg_other_rows * 0.7:
            is_header_row = True
    
    # Check for bold formatting in first row
    if table.rows[0]:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.bold:
                        is_header_row = True
                        break
    
    headers = table_data[0] if is_header_row and len(table_data) > 0 else []
    rows = table_data[1:] if is_header_row and len(table_data) > 1 else table_data
    
    if is_header_row and headers:
        # Format as header: value pairs
        for row in rows:
            for i, (header, value) in enumerate(zip(headers, row)):
                if value:  # Only add non-empty values
                    text = f"{header}{separator}{value}"
                    text_items.append(text)
    else:
        # Format as rows with cells separated by dashes
        for row in rows:
            if any(row):  # Only add rows with content
                row_text = ' - '.join(cell for cell in row if cell)
                text_items.append(row_text)
    
    return text_items

def tables_to_lists(docx_path, list_style='bullet', separator=': ', 
                    remove_tables=True, header_prefix="**"):
    """
    Convert all tables in a Word document to bullet/number lists.
    
    Args:
        docx_path: Path to the DOCX file
        list_style: 'bullet', 'number', or 'dash' (default: 'bullet')
        separator: String to separate headers from values (default: ': ')
        remove_tables: Whether to remove the original tables (default: True)
        header_prefix: Prefix for header text (default: "**")
    """
    docx_path = os.path.abspath(docx_path)

    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Error: Cannot find DOCX file at '{docx_path}'")

    doc = Document(docx_path)
    
    # Process tables in order
    tables = list(doc.tables)
    
    # Create list paragraphs with proper styles
    list_style_map = {
        'bullet': 'List Bullet',
        'number': 'List Number',
        'dash': 'List Bullet 2'
    }
    style_name = list_style_map.get(list_style, 'List Bullet')
    
    for table in tables:
        # Convert table to list text items
        text_items = table_to_list_text(table, list_style, separator)
        
        if not text_items:
            continue
        
        # Add list items after the table
        for text in text_items:
            doc.add_paragraph(text, style=style_name)
        
        # Remove the table if requested
        if remove_tables:
            # Access the table's element and remove it from its parent
            tbl = table._tbl
            parent = tbl.getparent()
            if parent is not None:
                parent.remove(tbl)
    
    doc.save(docx_path)
    print(f"Successfully converted {len(tables)} table(s) to lists in '{docx_path}'")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python tables_to_lists.py <docx_path> [list_style] [separator] [remove_tables]")
        print()
        print("Options:")
        print("  list_style: bullet, number, or dash (default: bullet)")
        print("  separator: string to separate headers from values (default: ': ')")
        print("  remove_tables: True or False (default: True)")
        print()
        print("Example: python tables_to_lists.py doc.docx bullet ': ' True")
        sys.exit(1)

    docx_path = sys.argv[1]
    list_style = sys.argv[2] if len(sys.argv) > 2 else 'bullet'
    separator = sys.argv[3] if len(sys.argv) > 3 else ': '
    remove_tables = sys.argv[4].lower() == 'true' if len(sys.argv) > 4 else True

    tables_to_lists(docx_path, list_style, separator, remove_tables)