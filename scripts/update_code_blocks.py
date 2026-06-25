#!/usr/bin/env python3
"""
Update code blocks in Word documents to use fixed-width font.

This script ensures that code blocks in the converted DOCX file use
a fixed-width (monospace) font for better readability.

Usage:
    python update_code_blocks.py <docx_path>
"""

import sys
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def update_code_blocks(docx_path):
    """
    Update code blocks to use fixed-width font.
    
    Args:
        docx_path: Path to the DOCX file
    """
    docx_path = os.path.abspath(docx_path)
    
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Error: Cannot find DOCX file at '{docx_path}'")
    
    doc = Document(docx_path)
    
    # Common fixed-width fonts to try (in order of preference)
    fixed_width_fonts = ['Courier New', 'Consolas', 'Monaco', 'Lucida Console', 'DejaVu Sans Mono']
    code_font = None
    
    # Check which fonts are available (try to find one that exists)
    # For simplicity, we'll use 'Courier New' as the default as it's widely available
    code_font = 'Courier New'
    
    code_block_count = 0
    
    for paragraph in doc.paragraphs:
        # Check if this paragraph contains code (has the 'Code' style or is in a code block)
        if paragraph.style.name.startswith('Code') or \
           paragraph.style.name == 'VerbatimChar' or \
           paragraph.style.name == 'SourceCode':
            
            # Set fixed-width font for all runs
            for run in paragraph.runs:
                run.font.name = code_font
                run.font.size = Pt(10)  # Slightly smaller font for code
            code_block_count += 1
    
    # Also check for code in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.style.name.startswith('Code') or \
                       paragraph.style.name == 'VerbatimChar' or \
                       paragraph.style.name == 'SourceCode':
                        
                        for run in paragraph.runs:
                            run.font.name = code_font
                            run.font.size = Pt(10)
                        code_block_count += 1
    
    doc.save(docx_path)
    print(f"Updated {code_block_count} code block(s) with fixed-width font '{code_font}'")
    return code_block_count


def main():
    if len(sys.argv) < 2:
        print("Usage: python update_code_blocks.py <docx_path>")
        print()
        print("Updates code blocks in a Word document to use fixed-width font.")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    try:
        count = update_code_blocks(docx_path)
        sys.exit(0)
    except Exception as e:
        print(f"❌ Error updating code blocks: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
