import sys
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def update_header(docx_path, title_text, classification=None, font_size=14, bold=True, align='center', classification_font_size=10):
    """
    Update all headers in a Word document with the given title text and classification.

    Args:
        docx_path: Path to the DOCX file
        title_text: Text to insert in the header
        classification: Classification text (e.g., 'UNCLASSIFIED') (default: None)
        font_size: Font size for the header text (default: 14)
        bold: Whether to make the text bold (default: True)
        align: Text alignment - 'left', 'center', or 'right' (default: 'center')
        classification_font_size: Font size for classification text (default: 10)
    """
    # Ensure the docx_path is an absolute path
    docx_path = os.path.abspath(docx_path)

    # Check if the file exists
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Error: Cannot find DOCX file at '{docx_path}'")

    # Load the document
    doc = Document(docx_path)

    # Map alignment strings to constants
    alignment_map = {
        'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
        'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
        'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
    }
    alignment = alignment_map.get(align.lower(), WD_PARAGRAPH_ALIGNMENT.CENTER)

    # Update the headers in all sections
    for section in doc.sections:
        header = section.header

        # Get all paragraphs first to avoid iteration issues
        paragraphs = list(header.paragraphs)
        
        # Clear all existing paragraphs by setting their text to empty
        for paragraph in paragraphs:
            paragraph.text = ''

        # Convert title to ALL CAPS if classification is present (military style)
        display_title = title_text.upper() if classification else title_text
        
        # Use or create first paragraph for the title
        if paragraphs:
            header_para = paragraphs[0]
        else:
            header_para = header.add_paragraph()
        
        # Set title text (ALL CAPS if classification exists)
        header_para.text = display_title

        # Apply styling to title
        for run in header_para.runs:
            run.font.size = Pt(font_size)
            run.font.bold = bold
        header_para.alignment = alignment
        
        # Add classification text on the RIGHT side of the same line if provided
        if classification:
            # Create a single line: TITLE                              UNCLASSIFIED
            # We'll use a tab or spaces to separate them
            header_para.text = f"{display_title}\t\t{classification}"
            
            # Style the entire line
            for run in header_para.runs:
                run.font.size = Pt(font_size)
                run.font.bold = bold

    # Save the updated document
    doc.save(docx_path)
    print(f"Successfully updated header in '{docx_path}' with title: '{title_text}'")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python update_header.py <docx_path> <title_text> [classification] [font_size] [bold] [align] [classification_font_size]")
        print("Example: python update_header.py doc.docx 'My Document Title' 'UNCLASSIFIED' 16 True center 10")
        sys.exit(1)

    docx_path = sys.argv[1]
    title_text = sys.argv[2]

    # Parse optional arguments with defaults
    classification = sys.argv[3] if len(sys.argv) > 3 else None
    font_size = int(sys.argv[4]) if len(sys.argv) > 4 else 14
    bold = sys.argv[5].lower() == 'true' if len(sys.argv) > 5 else True
    align = sys.argv[6] if len(sys.argv) > 6 else 'center'
    classification_font_size = int(sys.argv[7]) if len(sys.argv) > 7 else 10

    update_header(docx_path, title_text, classification, font_size, bold, align, classification_font_size)
