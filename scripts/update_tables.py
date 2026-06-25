import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

def _set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Check each border side
    for side in ['top', 'start', 'bottom', 'end', 'insideH', 'insideV']:
        border_name = f'{side}Border'
        border_val = kwargs.get(side, kwargs.get('all', None))
        if border_val is not None:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), border_val['style'])
            border.set(qn('w:sz'), str(border_val['size']))
            border.set(qn('w:space'), '0')
            if 'color' in border_val:
                border.set(qn('w:color'), border_val['color'])
            tcPr.append(border)

def _set_table_border(table, **kwargs):
    """Set border for all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, **kwargs)

def update_tables(docx_path, style='grid', border_size=6, border_color='auto', 
                  cell_padding=4, text_alignment='center', header_row=True,
                  header_bold=True, header_bg_color=None, alternating_rows=False,
                  font_size=10):
    """
    Update all tables in a Word document with better formatting.
    
    Args:
        docx_path: Path to the DOCX file
        style: Table style - 'grid', 'clean', 'borderless', or 'custom' (default: 'grid')
        border_size: Border size in points (default: 6)
        border_color: Border color as hex string (e.g., '4472C4') or 'auto' (default: 'auto')
        cell_padding: Cell padding in points (default: 4)
        text_alignment: Text alignment in cells - 'left', 'center', 'right' (default: 'center')
        header_row: Whether first row is a header (default: True)
        header_bold: Make header text bold (default: True)
        header_bg_color: Header background color as hex string (default: None)
        alternating_rows: Apply alternating row colors (default: False)
        font_size: Font size in points for table text (default: 10)
    """
    docx_path = os.path.abspath(docx_path)

    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Error: Cannot find DOCX file at '{docx_path}'")

    doc = Document(docx_path)

    # Define border styles
    border_styles = {
        'grid': {
            'all': {'style': 'single', 'size': border_size, 'color': border_color}
        },
        'clean': {
            'top': {'style': 'single', 'size': border_size, 'color': border_color},
            'bottom': {'style': 'single', 'size': border_size, 'color': border_color},
            'start': None,
            'end': None,
            'insideH': {'style': 'single', 'size': border_size//2, 'color': border_color},
            'insideV': None
        },
        'borderless': {},
        'custom': {
            'all': {'style': 'single', 'size': border_size, 'color': border_color}
        }
    }
    
    border_config = border_styles.get(style, border_styles['grid'])
    
    # Color mapping
    if border_color == 'auto':
        border_color = '4472C4'  # Default accent color
    if header_bg_color == 'auto':
        header_bg_color = 'D9E2F3'  # Light blue

    for table in doc.tables:
        # Enable autofit
        table.autofit = True
        
        # Set table alignment
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths to auto
        for col in table.columns:
            col.width = None
        
        # Apply borders based on style
        if border_config:
            _set_table_border(table, **border_config)
        
        # Apply cell padding
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.left_indent = Pt(cell_padding)
                    paragraph.paragraph_format.right_indent = Pt(cell_padding)
                    
                    # Text alignment
                    if text_alignment == 'left':
                        paragraph.alignment = 0  # WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif text_alignment == 'right':
                        paragraph.alignment = 2  # WD_PARAGRAPH_ALIGNMENT.RIGHT
                    else:
                        paragraph.alignment = 1  # WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Vertical alignment
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    # Apply font size to all runs in the cell
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)
                    
                    # Header row formatting
                    if header_row and row == table.rows[0]:
                        for run in paragraph.runs:
                            run.font.bold = header_bold
                        
                        if header_bg_color:
                            # Set cell shading
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            shade = OxmlElement('w:shd')
                            shade.set(qn('w:val'), 'clear')
                            shade.set(qn('w:color'), 'auto')
                            shade.set(qn('w:fill'), header_bg_color)
                            tcPr.append(shade)
                    
                    # Alternating row colors
                    if alternating_rows and row != table.rows[0]:
                        if row.index % 2 == 0:  # Even rows
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            shade = OxmlElement('w:shd')
                            shade.set(qn('w:val'), 'clear')
                            shade.set(qn('w:color'), 'auto')
                            shade.set(qn('w:fill'), 'F2F2F2')  # Light gray
                            tcPr.append(shade)

    doc.save(docx_path)
    print(f"Successfully updated {len(doc.tables)} table(s) in '{docx_path}' with style: '{style}'")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python update_tables.py <docx_path> [style] [border_size] [border_color] [cell_padding] [text_alignment]")
        print("       [header_row] [header_bold] [header_bg_color] [alternating_rows]")
        print()
        print("Styles: grid, clean, borderless, custom")
        print("Example: python update_tables.py doc.docx grid 8 4472C4 6 center True True D9E2F3 False")
        sys.exit(1)

    docx_path = sys.argv[1]
    style = sys.argv[2] if len(sys.argv) > 2 else 'grid'
    border_size = int(sys.argv[3]) if len(sys.argv) > 3 else 6
    border_color = sys.argv[4] if len(sys.argv) > 4 else 'auto'
    cell_padding = int(sys.argv[5]) if len(sys.argv) > 5 else 4
    text_alignment = sys.argv[6] if len(sys.argv) > 6 else 'center'
    header_row = sys.argv[7].lower() == 'true' if len(sys.argv) > 7 else True
    header_bold = sys.argv[8].lower() == 'true' if len(sys.argv) > 8 else True
    header_bg_color = sys.argv[9] if len(sys.argv) > 9 else None
    alternating_rows = sys.argv[10].lower() == 'true' if len(sys.argv) > 10 else False

    update_tables(docx_path, style, border_size, border_color, cell_padding, 
                  text_alignment, header_row, header_bold, header_bg_color, alternating_rows)
